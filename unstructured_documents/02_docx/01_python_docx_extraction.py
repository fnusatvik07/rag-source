"""
DOCX extraction using python-docx — the most feature-rich pure-Python approach.

python-docx gives you access to paragraphs, runs (inline formatting), styles,
tables, images, headers/footers, and more.  This script demonstrates:

  1. Paragraph extraction with style metadata
  2. Table extraction with cell data
  3. Heading-based document structure (hierarchy)
  4. Chunking by headings for RAG

Run:
    uv run python unstructured_documents/02_docx/01_python_docx_extraction.py
"""

import sys
from pathlib import Path

from docx import Document

# ---------------------------------------------------------------------------
# Make shared utilities importable from the project root
# ---------------------------------------------------------------------------
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))
from unstructured_documents.shared.chunking import (
    chunk_by_headings,
    preview_chunks,
)

SAMPLE_DOC = Path(__file__).resolve().parent / "sample_docs" / "simple_document.docx"
TABLES_DOC = Path(__file__).resolve().parent / "sample_docs" / "tables_document.docx"


# ===================================================================
# 1. Extract paragraphs with style information
# ===================================================================


def extract_paragraphs(doc_path: Path) -> list[dict]:
    """
    Extract every paragraph together with its style name and full text.

    python-docx exposes each block-level element as a Paragraph object.
    The `style.name` attribute gives you the Word style (e.g. "Heading 1",
    "Normal", "List Bullet") which is extremely valuable for downstream
    processing.
    """
    doc = Document(str(doc_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "style": para.style.name,
                "text": text,
                # Run-level detail: capture bold/italic spans
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                    }
                    for run in para.runs
                    if run.text.strip()
                ],
            }
        )
    return paragraphs


# ===================================================================
# 2. Extract tables
# ===================================================================


def extract_tables(doc_path: Path) -> list[list[list[str]]]:
    """
    Extract all tables in the document.

    Each table is returned as a list of rows, where each row is a list
    of cell text strings.  The first row is typically the header.
    """
    doc = Document(str(doc_path))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


# ===================================================================
# 3. Build a heading-based document hierarchy
# ===================================================================


def build_heading_hierarchy(doc_path: Path) -> list[dict]:
    """
    Walk through the document and group content under headings.

    This produces a flat list of sections, each with:
      - heading_level (int)  — 0 for non-heading preamble, 1-9 for headings
      - heading_text  (str)  — the heading text
      - body_parts    (list) — paragraphs and tables that follow the heading

    This is the most natural structure for RAG chunking: each section
    becomes a chunk (or is further split if too large).
    """
    doc = Document(str(doc_path))
    sections: list[dict] = []
    current_section: dict = {
        "heading_level": 0,
        "heading_text": "Preamble",
        "body_parts": [],
    }

    for para in doc.paragraphs:
        style_name = para.style.name  # e.g. "Heading 1", "Normal"
        text = para.text.strip()
        if not text:
            continue

        # Detect heading paragraphs
        if style_name.startswith("Heading"):
            # Save previous section if it has content
            if current_section["body_parts"]:
                sections.append(current_section)
            level = int(style_name.split()[-1])
            current_section = {
                "heading_level": level,
                "heading_text": text,
                "body_parts": [],
            }
        else:
            current_section["body_parts"].append(
                {
                    "type": "paragraph",
                    "style": style_name,
                    "text": text,
                }
            )

    # Don't forget the last section
    if current_section["body_parts"]:
        sections.append(current_section)

    return sections


# ===================================================================
# 4. Convert to markdown text (for heading-based chunking)
# ===================================================================


def docx_to_markdown(doc_path: Path) -> str:
    """
    Convert the DOCX to a simple markdown string so we can reuse
    the shared heading-aware chunker.

    Mapping:
      Heading 1  ->  # Heading
      Heading 2  ->  ## Heading
      Heading 3  ->  ### Heading
      List Bullet -> - text
      List Number -> 1. text
      Normal      ->  text (blank-line separated)
    """
    doc = Document(str(doc_path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name

        if style.startswith("Heading"):
            level = int(style.split()[-1])
            prefix = "#" * level
            lines.append(f"\n{prefix} {text}\n")
        elif style == "List Bullet":
            lines.append(f"- {text}")
        elif style.startswith("List Number"):
            lines.append(f"1. {text}")
        else:
            lines.append(f"\n{text}\n")

    return "\n".join(lines)


# ===================================================================
# Main demonstration
# ===================================================================


def main() -> None:
    print("=" * 70)
    print("DOCX Extraction with python-docx")
    print("=" * 70)

    # ------------------------------------------------------------------
    # 1. Paragraphs + styles
    # ------------------------------------------------------------------
    print("\n\n--- 1. Paragraph Extraction (with styles) ---\n")
    paragraphs = extract_paragraphs(SAMPLE_DOC)
    print(f"Total non-empty paragraphs: {len(paragraphs)}\n")

    for p in paragraphs[:8]:
        style_tag = f"[{p['style']}]"
        snippet = p["text"][:90] + ("..." if len(p["text"]) > 90 else "")
        print(f"  {style_tag:20s}  {snippet}")

        # Show bold/italic runs if present
        rich_runs = [r for r in p["runs"] if r["bold"] or r["italic"]]
        for r in rich_runs:
            flags = []
            if r["bold"]:
                flags.append("BOLD")
            if r["italic"]:
                flags.append("ITALIC")
            print(f'  {"":20s}    ^ {", ".join(flags)}: "{r["text"][:60]}"')

    if len(paragraphs) > 8:
        print(f"\n  ... and {len(paragraphs) - 8} more paragraphs")

    # ------------------------------------------------------------------
    # 2. Tables
    # ------------------------------------------------------------------
    print("\n\n--- 2. Table Extraction ---\n")
    tables = extract_tables(TABLES_DOC)
    print(f"Total tables found: {len(tables)}\n")

    for idx, table in enumerate(tables):
        print(f"  Table {idx + 1}: {len(table)} rows x {len(table[0])} cols")
        # Show header row
        print(f"    Header: {table[0]}")
        # Show first data row
        if len(table) > 1:
            print(f"    Row  1: {table[1]}")
        print()

    # ------------------------------------------------------------------
    # 3. Heading hierarchy
    # ------------------------------------------------------------------
    print("\n--- 3. Heading-Based Document Structure ---\n")
    sections = build_heading_hierarchy(SAMPLE_DOC)
    print(f"Total sections: {len(sections)}\n")

    for sec in sections:
        indent = "  " * sec["heading_level"]
        n_parts = len(sec["body_parts"])
        print(
            f"  {indent}H{sec['heading_level']}: {sec['heading_text']}  "
            f"({n_parts} body element{'s' if n_parts != 1 else ''})"
        )

    # ------------------------------------------------------------------
    # 4. Chunking by headings (via markdown conversion)
    # ------------------------------------------------------------------
    print("\n\n--- 4. Heading-Aware Chunking for RAG ---\n")
    md_text = docx_to_markdown(SAMPLE_DOC)

    # Show a snippet of the markdown conversion
    print("Markdown conversion (first 500 chars):")
    print("-" * 40)
    print(md_text[:500])
    print("-" * 40)

    # Use the shared heading chunker
    chunks = chunk_by_headings(md_text)
    preview_chunks(chunks, max_preview=4, max_chars=250)


if __name__ == "__main__":
    main()
