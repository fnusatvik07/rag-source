"""
Generate sample DOCX files for demonstrating different extraction methods.

Creates three documents with varying complexity:
  1. simple_document.docx  - headings, paragraphs, bullet points, bold/italic
  2. tables_document.docx  - multiple tables with surrounding text
  3. styled_document.docx  - wide variety of Word styles and numbered lists

Run:
    uv run python unstructured_documents/02_docx/sample_docs/generate_samples.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

SAMPLE_DIR = Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# 1. simple_document.docx
# ---------------------------------------------------------------------------


def create_simple_document() -> None:
    """Create a document with headings, paragraphs, bullet points, and rich text."""

    doc = Document()

    # --- Title / H1 ---
    doc.add_heading("Understanding Climate Change: Causes, Effects, and Solutions", level=1)

    doc.add_paragraph(
        "Climate change is one of the most pressing challenges facing humanity in the "
        "twenty-first century. This document provides an overview of the science behind "
        "climate change, its observable effects around the world, and the range of "
        "strategies being pursued to mitigate and adapt to a warming planet."
    )

    # --- H2: The Science ---
    doc.add_heading("The Science of Climate Change", level=2)

    para = doc.add_paragraph(
        "Earth's climate is determined by the balance between incoming solar radiation "
        "and outgoing infrared radiation. "
    )
    run = para.add_run("Greenhouse gases")
    run.bold = True
    para.add_run(
        " such as carbon dioxide (CO₂), methane (CH₄), and nitrous oxide (N₂O) trap "
        "a portion of the outgoing radiation, warming the lower atmosphere. This natural "
        "process is essential for life, maintaining global average temperatures roughly "
        "33 °C warmer than they would otherwise be."
    )

    doc.add_paragraph(
        "Human activities — primarily the burning of fossil fuels, deforestation, and "
        "industrial agriculture — have increased atmospheric CO₂ concentrations from "
        "about 280 parts per million (ppm) in the pre-industrial era to over 420 ppm "
        "today. This enhanced greenhouse effect is the dominant driver of observed "
        "warming since the mid-twentieth century."
    )

    # --- H3: Key Indicators ---
    doc.add_heading("Key Climate Indicators", level=3)

    doc.add_paragraph("Scientists track several indicators to monitor the state of the climate system:")

    bullets = [
        "Global mean surface temperature has risen approximately 1.1 °C above pre-industrial levels as of 2023.",
        "Arctic sea-ice extent has declined by roughly 13 % per decade since satellite records began in 1979.",
        "Global mean sea level has risen about 20 cm since 1900, with the rate of rise accelerating in recent decades.",
        "Ocean heat content has increased steadily, with the upper 2,000 metres of "
        "the ocean absorbing over 90 % of the excess heat.",
        "Atmospheric methane concentrations have more than doubled since "
        "pre-industrial times, driven by agriculture, fossil-fuel extraction, and "
        "wetland feedbacks.",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    # --- H2: Effects ---
    doc.add_heading("Observable Effects of Climate Change", level=2)

    para = doc.add_paragraph("The effects of climate change are ")
    run = para.add_run("already being felt")
    run.italic = True
    para.add_run(" across every continent and ocean. Some of the most significant impacts include:")

    doc.add_heading("Extreme Weather Events", level=3)

    doc.add_paragraph(
        "Warmer temperatures increase evaporation and add energy to the atmosphere, "
        "intensifying storms, heatwaves, droughts, and heavy precipitation events. "
        "The frequency of Category 4 and 5 hurricanes has increased by about 25 % "
        "over the past four decades. Record-breaking heatwaves have struck every "
        "inhabited continent, with Western Europe experiencing temperatures above "
        "40 °C with increasing regularity."
    )

    doc.add_heading("Ecosystem Disruption", level=3)

    doc.add_paragraph(
        "Species are shifting their ranges poleward and to higher elevations, "
        "disrupting ecological communities that have been stable for millennia. "
        "Coral reefs have suffered multiple mass bleaching events; the Great Barrier "
        "Reef alone experienced five mass bleaching events between 2016 and 2024. "
        "Phenological mismatches — where the timing of events such as flowering and "
        "insect emergence falls out of sync — threaten pollination and food webs."
    )

    # --- H2: Solutions ---
    doc.add_heading("Strategies for Mitigation and Adaptation", level=2)

    doc.add_paragraph(
        "Addressing climate change requires a two-pronged approach: mitigation, which "
        "aims to reduce emissions and enhance carbon sinks, and adaptation, which "
        "prepares societies for the changes that are already unavoidable."
    )

    doc.add_heading("Mitigation Strategies", level=3)

    mitigation_points = [
        "Transitioning electricity generation from fossil fuels to renewables such as "
        "solar, wind, and geothermal energy.",
        "Electrifying transportation and heating to take advantage of cleaner grids.",
        "Improving energy efficiency in buildings, industry, and agriculture.",
        "Protecting and restoring forests, wetlands, and other natural carbon sinks.",
        "Developing and deploying carbon capture, utilisation, and storage (CCUS) "
        "technologies for hard-to-abate sectors.",
    ]
    for point in mitigation_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Adaptation Strategies", level=3)

    adaptation_points = [
        "Building resilient infrastructure designed for future climate conditions.",
        "Developing drought-resistant crop varieties and sustainable water management.",
        "Strengthening early-warning systems for extreme weather events.",
        "Implementing nature-based solutions such as mangrove restoration for coastal protection.",
    ]
    for point in adaptation_points:
        doc.add_paragraph(point, style="List Bullet")

    # --- H2: Conclusion ---
    doc.add_heading("Conclusion", level=2)

    para = doc.add_paragraph(
        "Climate change is a defining issue of our time. The scientific evidence is "
        "unequivocal: human activities have warmed the planet, and the impacts are "
        "accelerating. However, pathways to limit warming to 1.5 °C remain available "
        "if rapid, far-reaching action is taken across all sectors of the economy. "
    )
    run = para.add_run(
        "The choices made in the next decade will determine the trajectory of the climate for centuries to come."
    )
    run.bold = True
    run.italic = True

    path = SAMPLE_DIR / "simple_document.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------
# 2. tables_document.docx
# ---------------------------------------------------------------------------


def create_tables_document() -> None:
    """Create a document with multiple tables and explanatory text."""

    doc = Document()

    doc.add_heading("Quarterly Business Report — Q3 2025", level=1)

    doc.add_paragraph(
        "This report summarises the financial performance, inventory position, and "
        "staffing overview for Acme Corp during the third quarter of 2025. All figures "
        "are presented in US dollars unless otherwise noted."
    )

    # --- Table 1: Financial Summary ---
    doc.add_heading("Financial Summary", level=2)

    doc.add_paragraph(
        "The following table provides a high-level view of revenue, cost of goods "
        "sold (COGS), gross profit, operating expenses, and net income for Q3 2025 "
        "compared with the prior quarter and the same quarter last year."
    )

    financial_headers = ["Metric", "Q3 2025", "Q2 2025", "Q3 2024", "YoY Change"]
    financial_data = [
        ["Revenue", "$12,450,000", "$11,820,000", "$10,900,000", "+14.2 %"],
        ["Cost of Goods Sold", "$7,470,000", "$7,210,000", "$6,760,000", "+10.5 %"],
        ["Gross Profit", "$4,980,000", "$4,610,000", "$4,140,000", "+20.3 %"],
        ["Operating Expenses", "$2,850,000", "$2,790,000", "$2,610,000", "+9.2 %"],
        ["Net Income", "$2,130,000", "$1,820,000", "$1,530,000", "+39.2 %"],
    ]

    table = doc.add_table(rows=1, cols=len(financial_headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(financial_headers):
        table.rows[0].cells[i].text = header
    for row_data in financial_data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value

    doc.add_paragraph(
        "Revenue grew 14.2 % year-over-year, driven primarily by strong demand in the "
        "North American and European markets. Gross margin improved to 40.0 %, up from "
        "38.0 % in Q3 2024, reflecting operational efficiencies and favourable input costs."
    )

    # --- Table 2: Inventory Status ---
    doc.add_heading("Inventory Status", level=2)

    doc.add_paragraph(
        "Effective inventory management remains a priority. The table below shows "
        "current stock levels across our main product categories."
    )

    inventory_headers = [
        "Product Category",
        "SKU Count",
        "Units in Stock",
        "Reorder Point",
        "Status",
    ]
    inventory_data = [
        ["Electronics", "142", "34,500", "10,000", "Adequate"],
        ["Home Appliances", "87", "12,200", "5,000", "Adequate"],
        ["Office Supplies", "215", "98,000", "25,000", "Adequate"],
        ["Industrial Tools", "63", "4,800", "5,000", "Low — reorder initiated"],
        ["Automotive Parts", "178", "22,100", "8,000", "Adequate"],
        ["Health & Safety", "54", "7,300", "7,000", "Marginal"],
    ]

    table = doc.add_table(rows=1, cols=len(inventory_headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(inventory_headers):
        table.rows[0].cells[i].text = header
    for row_data in inventory_data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value

    doc.add_paragraph(
        "Industrial Tools inventory dipped below the reorder point in late August; "
        "a purchase order for 3,000 additional units was placed on September 5. "
        "Health & Safety stock is marginal and will be reviewed in the upcoming "
        "planning cycle."
    )

    # --- Table 3: Employee Directory ---
    doc.add_heading("Key Personnel", level=2)

    doc.add_paragraph(
        "The following directory lists the leadership team responsible for Q3 "
        "operations and their primary contact information."
    )

    emp_headers = ["Name", "Title", "Department", "Office", "Email"]
    emp_data = [
        ["Sarah Chen", "CEO", "Executive", "New York", "s.chen@acme.com"],
        ["James Okafor", "CFO", "Finance", "New York", "j.okafor@acme.com"],
        [
            "Maria Gonzalez",
            "VP of Engineering",
            "Engineering",
            "San Francisco",
            "m.gonzalez@acme.com",
        ],
        ["David Kim", "VP of Sales", "Sales", "Chicago", "d.kim@acme.com"],
        ["Priya Patel", "VP of Operations", "Operations", "Dallas", "p.patel@acme.com"],
        ["Thomas Weber", "Head of HR", "Human Resources", "London", "t.weber@acme.com"],
        [
            "Aisha Mohammed",
            "General Counsel",
            "Legal",
            "New York",
            "a.mohammed@acme.com",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(emp_headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(emp_headers):
        table.rows[0].cells[i].text = header
    for row_data in emp_data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value

    doc.add_heading("Outlook", level=2)

    doc.add_paragraph(
        "Management expects revenue growth to continue in Q4, supported by the "
        "holiday season and new product launches planned for November. Capital "
        "expenditure for the quarter is budgeted at $1.2 million, primarily for "
        "warehouse automation and IT infrastructure upgrades."
    )

    path = SAMPLE_DIR / "tables_document.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------
# 3. styled_document.docx
# ---------------------------------------------------------------------------


def create_styled_document() -> None:
    """Create a document exercising many built-in Word styles."""

    doc = Document()

    # Title & subtitle
    doc.add_paragraph("The Art of Software Architecture", style="Title")
    doc.add_paragraph("A Practical Guide to Designing Maintainable Systems", style="Subtitle")

    # --- Heading 1 ---
    doc.add_heading("Introduction", level=1)

    doc.add_paragraph(
        "Software architecture is the set of high-level decisions that shape a "
        "system's structure, behaviour, and quality attributes. Good architecture "
        "makes a system easier to understand, develop, test, deploy, and operate. "
        "Poor architecture leads to rigid codebases where even small changes "
        "cascade unpredictably through the system."
    )

    doc.add_paragraph(
        "This guide distils decades of collective industry experience into concise, "
        "actionable advice. It is intended for developers who are transitioning into "
        "architecture roles, as well as seasoned architects looking for a refresher.",
        style="Normal",
    )

    # --- Quote style ---
    doc.add_paragraph(
        '"Architecture is the decisions that you wish you could get right early in '
        "a project, but that you are not necessarily more likely to get them right "
        'than any other." — Ralph Johnson',
        style="Quote",
    )

    # --- Heading 1 ---
    doc.add_heading("Fundamental Principles", level=1)

    doc.add_heading("Separation of Concerns", level=2)

    doc.add_paragraph(
        "Each module or component should address a single concern. When concerns are "
        "well-separated, changes to one part of the system have minimal impact on "
        "others. This principle underlies patterns such as MVC, hexagonal architecture, "
        "and micro-services."
    )

    doc.add_heading("Loose Coupling and High Cohesion", level=2)

    doc.add_paragraph(
        "Coupling measures how much one module depends on another. Cohesion measures "
        "how strongly the elements within a module belong together. Aim for low "
        "coupling between modules and high cohesion within them. This makes the system "
        "easier to test, refactor, and extend."
    )

    doc.add_heading("SOLID Principles", level=2)

    doc.add_paragraph("The SOLID principles provide a foundation for object-oriented design:")

    # Numbered list
    numbered_items = [
        "Single Responsibility Principle — a class should have one, and only one, reason to change.",
        "Open/Closed Principle — software entities should be open for extension but closed for modification.",
        "Liskov Substitution Principle — subtypes must be substitutable for their "
        "base types without altering program correctness.",
        "Interface Segregation Principle — clients should not be forced to depend on interfaces they do not use.",
        "Dependency Inversion Principle — high-level modules should not depend on "
        "low-level modules; both should depend on abstractions.",
    ]
    for item in numbered_items:
        doc.add_paragraph(item, style="List Number")

    # --- Heading 1 ---
    doc.add_heading("Common Architectural Patterns", level=1)

    doc.add_heading("Layered Architecture", level=2)

    doc.add_paragraph(
        "The layered (or n-tier) architecture organises code into horizontal layers, "
        "typically presentation, business logic, and data access. Each layer "
        "communicates only with the layer directly below it. This pattern is simple "
        "and widely understood, making it a good default for many applications."
    )

    doc.add_heading("Event-Driven Architecture", level=2)

    doc.add_paragraph(
        "In event-driven systems, components communicate by producing and consuming "
        "events rather than making direct calls. This decouples producers from "
        "consumers and supports asynchronous processing, replay, and auditing. "
        "Popular implementations include Apache Kafka, RabbitMQ, and cloud-native "
        "services like AWS EventBridge."
    )

    doc.add_heading("Micro-services", level=2)

    doc.add_paragraph(
        "Micro-services decompose a system into independently deployable services, "
        "each owning its data and exposing a well-defined API. Benefits include "
        "independent scaling, technology diversity, and team autonomy. Challenges "
        "include distributed-system complexity, network latency, and operational "
        "overhead."
    )

    # --- Code-style text ---
    doc.add_heading("Code Example: Dependency Injection", level=2)

    doc.add_paragraph("The following pseudo-code shows constructor-based dependency injection:")

    # Use a monospace font for code
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        "class OrderService:\n"
        "    def __init__(self, repository: OrderRepository,\n"
        "                 notifier: NotificationService):\n"
        "        self._repository = repository\n"
        "        self._notifier = notifier\n"
        "\n"
        "    def place_order(self, order: Order) -> None:\n"
        "        self._repository.save(order)\n"
        "        self._notifier.send_confirmation(order)\n"
    )
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph(
        "By injecting dependencies through the constructor, the OrderService "
        "can be tested with mock implementations and is not coupled to any "
        "concrete persistence or notification technology."
    )

    # --- Another Quote ---
    doc.add_paragraph('"Make each program do one thing well." — Unix Philosophy', style="Quote")

    # --- Heading 1 ---
    doc.add_heading("Quality Attributes", level=1)

    doc.add_paragraph(
        "Architecture decisions are driven by quality attributes — the non-functional "
        "requirements that describe how the system should behave under various conditions."
    )

    quality_attrs = [
        "Performance — response time, throughput, and resource utilisation under load.",
        "Scalability — the ability to handle increased load by adding resources.",
        "Availability — the proportion of time the system is operational and accessible.",
        "Security — protection against unauthorised access, data breaches, and attacks.",
        "Maintainability — how easily the system can be modified, extended, and debugged.",
        "Testability — how easily the system's components can be verified in isolation.",
    ]
    for attr in quality_attrs:
        doc.add_paragraph(attr, style="List Number")

    # --- Conclusion ---
    doc.add_heading("Conclusion", level=1)

    doc.add_paragraph(
        "Software architecture is not a one-time activity but an ongoing discipline. "
        "As requirements evolve and technologies change, the architecture must adapt. "
        "By grounding decisions in proven principles and patterns, and by continuously "
        "evaluating quality attributes, teams can build systems that endure."
    )

    doc.add_paragraph(
        '"The best architectures are grown, not designed." — Adapted from Fred Brooks',
        style="Quote",
    )

    path = SAMPLE_DIR / "styled_document.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating sample DOCX files ...\n")
    create_simple_document()
    create_tables_document()
    create_styled_document()
    print("\nDone. All sample documents are in:", SAMPLE_DIR)
